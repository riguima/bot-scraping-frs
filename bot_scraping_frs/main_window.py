import asyncio
from io import BytesIO

import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Pt
from openpyxl import load_workbook
from openpyxl.drawing.image import Image
from openpyxl.styles import Alignment, Font, PatternFill
from PySide6 import QtCore, QtWidgets
from sqlalchemy import select

from bot_scraping_frs.browser import get_all_pages_data
from bot_scraping_frs.database import Session
from bot_scraping_frs.models import Product
from bot_scraping_frs.utils import (convert_value, format_number,
                                    get_all_images_content)


class MainWindow(QtWidgets.QWidget):
    def __init__(self):
        super().__init__()
        self.setWindowTitle('Bot Scraping FRS')
        self.setFixedSize(400, 250)

        with open('styles.qss', 'r') as f:
            self.setStyleSheet(f.read())

        self.message_box = QtWidgets.QMessageBox()

        with Session() as session:
            products = session.scalars(select(Product)).all()
            self.products_label = QtWidgets.QLabel(
                f'Produtos: {len(products)}',
                alignment=QtCore.Qt.AlignmentFlag.AlignCenter,
            )
            self.products_label.setStyleSheet('font-weight: bold')

        self.url_label = QtWidgets.QLabel('URL')
        self.url_input = QtWidgets.QLineEdit()
        self.url_layout = QtWidgets.QHBoxLayout()
        self.url_layout.addWidget(self.url_label)
        self.url_layout.addWidget(self.url_input)

        self.run_button = QtWidgets.QPushButton('Rodar')
        self.run_button.clicked.connect(self.run)

        self.export_to_excel_button = QtWidgets.QPushButton(
            'Exportar para Excel'
        )
        self.export_to_excel_button.clicked.connect(self.export_to_excel)

        self.export_to_pdf_button = QtWidgets.QPushButton('Exportar para PDF')
        self.export_to_pdf_button.clicked.connect(self.export_to_pdf)

        self.main_layout = QtWidgets.QVBoxLayout(self)
        self.main_layout.addWidget(self.products_label)
        self.main_layout.addLayout(self.url_layout)
        self.main_layout.addWidget(self.run_button)
        self.main_layout.addWidget(self.export_to_excel_button)
        self.main_layout.addWidget(self.export_to_pdf_button)

    @QtCore.Slot()
    def run(self):
        if not self.url_input.text():
            self.message_box.setText('Preencha a URL')
            self.message_box.show()
        self.message_box.setText('Aguarde...')
        self.message_box.exec()
        with Session() as session:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            data = loop.run_until_complete(
                get_all_pages_data(self.url_input.text())
            )
            loop.close()
            for item in data:
                query = select(Product).where(Product.url == item['url'])
                item_model = session.scalars(query).first()
                if item_model:
                    item_model.foto = item['foto']
                    item_model.codigo = item['codigo']
                    item_model.descricao = item['descricao']
                    item_model.compra = item['valor']
                    item_model.venda = convert_value(item['valor'])
                    item_model.tamanhos = item['tamanhos']
                else:
                    item_model = Product(
                        url=item['url'],
                        foto=item['foto'],
                        codigo=item['codigo'],
                        descricao=item['descricao'],
                        compra=item['valor'],
                        venda=convert_value(item['valor']),
                        tamanhos=item['tamanhos'],
                    )
                    session.add(item_model)
            session.commit()
            self.update_products_label()
            self.url_input.setText('')
            self.message_box.setText('Finalizado!')
            self.message_box.show()

    def update_products_label(self):
        with Session() as session:
            products = session.scalars(select(Product)).all()
            self.products_label.setText(f'Produtos: {len(products)}')

    @QtCore.Slot()
    def export_to_excel(self):
        file_path = QtWidgets.QFileDialog.getSaveFileName()[0]
        if '.xlsx' not in file_path:
            file_path = file_path + '.xlsx'
        with Session() as session:
            items = []
            for item in session.scalars(select(Product)).all():
                items.append(
                    {
                        'Foto': item.foto,
                        'Código': item.codigo,
                        'Descrição': item.descricao,
                        'Compra': format_number(item.compra, '£'),
                        'Venda': format_number(item.venda, 'R$'),
                        'Tamanhos': item.tamanhos,
                    }
                )
            df = pd.DataFrame(items)
            df.to_excel(file_path, index=False, startrow=1)
            wb = load_workbook(file_path)
            ws = wb.active
            ws.merge_cells('A1:F1')
            ws['A1'] = 'Catálogo de produtos disponíveis'
            ws.column_dimensions['A'].width = 18
            images_contents = self.get_images_contents()
            for image_content, cell in zip(
                images_contents, ws[f'A3:A{len(images_contents) + 2}']
            ):
                cell[0].value = ''
                ws.row_dimensions[cell[0].row].height = 100
                image = Image(BytesIO(image_content))
                image.width = 120
                image.height = 120
                ws.add_image(image, f'A{cell[0].row}')
            for letter in ['A', 'B', 'C', 'D', 'E', 'F']:
                self.format_column_cells(ws, letter)
                if letter != 'A':
                    self.set_width_to_content(ws, letter)
            wb.save(file_path)
        self.message_box.setText('Finalizado!')
        self.message_box.show()

    @QtCore.Slot()
    def export_to_pdf(self):
        file_path = QtWidgets.QFileDialog.getSaveFileName()[0]
        if '.docx' not in file_path:
            file_path = file_path + '.docx'
        with Session() as session:
            products = session.scalars(select(Product)).all()
            document = Document()
            section = document.sections[0]
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)
            section.left_margin = Cm(1)
            section.right_margin = Cm(1)
            paragraph = document.add_paragraph(
                'Catálogo de produtos disponíveis'
            )
            paragraph.text = 'Catálogo de produtos disponíveis'
            paragraph.runs[0].font.bold = True
            paragraph.runs[0].font.size = Pt(20)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            table = document.add_table(rows=len(products) + 1, cols=4)
            table.cell(0, 0).text = 'Foto'
            table.cell(0, 1).text = 'Código'
            table.cell(0, 2).text = 'Descrição'
            table.cell(0, 3).text = 'Valor'
            table.cell(0, 0)._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            )
            table.cell(0, 1)._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            )
            table.cell(0, 2)._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            )
            table.cell(0, 3)._tc.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="000000"/>'.format(nsdecls('w')))
            )
            images_contents = self.get_images_contents()
            for e, product in enumerate(products):
                run = table.cell(e + 1, 0).paragraphs[0].add_run()
                run.add_picture(BytesIO(images_contents[e]), width=Cm(4))
                table.cell(e + 1, 1).add_paragraph(product.codigo)
                table.cell(e + 1, 2).add_paragraph(product.descricao)
                table.cell(e + 1, 3).add_paragraph(
                    format_number(product.venda, 'R$')
                )
            self.make_rows_bold(table.rows[0])
            self.format_table_rows(table)
            document.save(file_path)
        self.message_box.setText('Finalizado!')
        self.message_box.show()

    def set_width_to_content(self, ws, column):
        max_length = 0
        for cell in ws[f'{column}:{column}']:
            if max_length < len(str(cell.value)):
                max_length = len(str(cell.value))
        ws.column_dimensions[column].width = max_length + 2

    def format_column_cells(self, ws, column):
        fill = PatternFill(
            fill_type='solid', start_color='000000', end_color='000000'
        )
        alignment = Alignment(horizontal='center', vertical='center')
        for cell in ws[f'{column}:{column}']:
            if cell.row == 1:
                cell.font = Font(size=20, bold=True)
                cell.alignment = alignment
                continue
            elif cell.row == 2:
                cell.font = Font(size=12, bold=True)
                cell.fill = fill
            cell.alignment = alignment
            cell.font = Font(size=12)

    def make_rows_bold(self, *rows):
        for row in rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    def format_table_rows(self, table):
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(12)

    def get_images_contents(self):
        with Session() as session:
            images_urls = [
                product.foto
                for product in session.scalars(select(Product)).all()
            ]
            loop = asyncio.new_event_loop()
            result = loop.run_until_complete(
                get_all_images_content(images_urls)
            )
            loop.close()
            return result
